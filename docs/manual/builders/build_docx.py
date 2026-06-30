import sys, os, shutil
HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(os.path.dirname(HERE))
sys.path.insert(0, ROOT)

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from docs.manual.builders.chapters_meta import CHAPTERS
from docs.manual.builders.markdown_parser import parse_chapter, PRINT_RE
from docs.manual.builders.markdown_parser import _clean as md_clean

CONTENT_DIR = os.path.join(os.path.dirname(HERE), "content")
OUT_BASE = r"C:\Users\vinic\Desktop\OmniMob - Padrao Diamond"
DIST_DIR = os.path.join(os.path.dirname(HERE), "dist", "docx")
os.makedirs(DIST_DIR, exist_ok=True)

GREEN_HEX = "2EC76A"
DARK_HEX = "0A0A0A"
TD = RGBColor(0x33, 0x33, 0x33)
GR = RGBColor(0x2E, 0xC7, 0x6A)
GD = RGBColor(0xD4, 0xAF, 0x37)
WH = RGBColor(0xFF, 0xFF, 0xFF)
LT = RGBColor(0xCC, 0xCC, 0xCC)
GY = RGBColor(0x71, 0x71, 0x7A)


def sh(cell, hx):
    tp = cell._tc.get_or_add_tcPr()
    e = OxmlElement("w:shd")
    e.set(qn("w:val"), "clear")
    e.set(qn("w:color"), "auto")
    e.set(qn("w:fill"), hx)
    tp.append(e)


def hr(para, hx):
    p = para._p
    pp_el = p.get_or_add_pPr()
    b = OxmlElement("w:pBdr")
    bt = OxmlElement("w:bottom")
    bt.set(qn("w:val"), "single")
    bt.set(qn("w:sz"), "12")
    bt.set(qn("w:space"), "1")
    bt.set(qn("w:color"), hx)
    b.append(bt)
    pp_el.append(b)


def fnt(run, sz=11, color=None, bold=False, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(sz)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def cover(doc, ch, alabel, acolor):
    tb = doc.add_table(rows=1, cols=1)
    tb.autofit = False
    c = tb.cell(0, 0)
    c.width = Inches(7)
    sh(c, DARK_HEX)
    p = c.paragraphs[0]
    fnt(p.add_run("OmniMob"), sz=24, bold=True, color=GR)
    fnt(p.add_run("   |   Manual Oficial Padrao Diamond"), sz=11, color=GD)
    p = c.add_paragraph()
    fnt(p.add_run("Modulo " + ("%02d" % ch["module_num"]) + " - " + ch["category"]), sz=13, bold=True, color=GR)
    p = c.add_paragraph()
    fnt(p.add_run(ch["title"]), sz=28, bold=True, color=WH)
    p = c.add_paragraph()
    fnt(p.add_run(ch["subtitle"]), sz=12, color=LT)
    p = c.add_paragraph()
    fnt(p.add_run(alabel.upper()), sz=12, bold=True, color=acolor)
    c.add_paragraph()
    doc.add_paragraph()


def hh(doc, t, lv=1):
    p = doc.add_paragraph()
    if lv == 1:
        fnt(p.add_run(t), sz=18, bold=True, color=GR)
        hr(p, GREEN_HEX)
    elif lv == 2:
        fnt(p.add_run(t), sz=14, bold=True, color=GD)
    else:
        fnt(p.add_run(t), sz=12, bold=True, color=TD)


def pp(doc, t, sz=11, color=None, bold=False, italic=False, align=None):
    if color is None:
        color = TD
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    fnt(p.add_run(t), sz=sz, color=color, bold=bold, italic=italic)
    return p


def bl(doc, t):
    p = doc.add_paragraph(style="List Bullet")
    fnt(p.add_run(t), sz=11, color=TD)


def pc(doc, pr):
    tb = doc.add_table(rows=1, cols=1)
    c = tb.cell(0, 0)
    sh(c, "1C1C1C")
    p = c.paragraphs[0]
    fnt(p.add_run("[ PRINT " + (pr["id"] or "") + " ]"), sz=11, bold=True, color=GD)
    p = c.add_paragraph()
    fnt(p.add_run("Tela: "), sz=10, bold=True, color=GR)
    fnt(p.add_run(pr["screen"]), sz=10, color=WH)
    if pr.get("logged_as"):
        p = c.add_paragraph()
        fnt(p.add_run("Logado como: "), sz=10, bold=True, color=GR)
        fnt(p.add_run(pr["logged_as"]), sz=10, color=LT)
    if pr.get("notes"):
        p = c.add_paragraph()
        fnt(p.add_run("Anotacoes:"), sz=10, bold=True, color=GR)
        for i, n in enumerate(pr["notes"], 1):
            p = c.add_paragraph()
            fnt(p.add_run("  " + str(i) + ". " + n), sz=10, color=LT)
    doc.add_paragraph()


def is_sep(s):
    if not s:
        return False
    for ch in s:
        if not (ch == "-" or ch == ":" or ch.isspace()):
            return False
    return True


def is_bul(s):
    if len(s) < 2:
        return False
    if s[0] in ("-", "*") and s[1] == " ":
        return True
    j = 0
    while j < len(s) and s[j].isdigit():
        j += 1
    if j > 0 and j + 1 < len(s) and s[j] == "." and s[j + 1] == " ":
        return True
    return False


def strip_bul(s):
    if s[0] in ("-", "*"):
        return s[2:].strip()
    j = 0
    while j < len(s) and s[j].isdigit():
        j += 1
    return s[j + 2:].strip()


def render(doc, body):
    in_code = False
    rows = []
    in_t = [False]

    def flush():
        if not rows:
            return
        cols = max(len(r) for r in rows)
        tb = doc.add_table(rows=len(rows), cols=cols)
        tb.style = "Light Grid Accent 1"
        for ri, row in enumerate(rows):
            for ci, ct in enumerate(row):
                if ci < cols:
                    cc = tb.cell(ri, ci)
                    cc.text = ""
                    p = cc.paragraphs[0]
                    fnt(p.add_run(ct.strip()), sz=10, bold=(ri == 0),
                        color=(GR if ri == 0 else TD))
                    if ri == 0:
                        sh(cc, "0A0A0A")
        doc.add_paragraph()
        in_t[0] = False
        rows.clear()

    for raw in body.split("\n"):
        ln = raw.rstrip()
        st = ln.strip()
        if st.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        if st.startswith("|") and st.endswith("|"):
            cs = [c.strip() for c in st.split("|")[1:-1]]
            if cs and all(is_sep(c) for c in cs):
                continue
            rows.append(cs)
            in_t[0] = True
            continue
        elif in_t[0]:
            flush()
        if not st:
            continue
        m = PRINT_RE.search(ln)
        if m:
            pid, sc, lg, nr = m.groups()
            pr = {"id": pid.strip() if pid else "", "screen": (sc or "").strip(),
                  "logged_as": (lg or "").strip(), "notes": []}
            if nr:
                for n in nr.split(","):
                    nn = n.strip()
                    if nn:
                        pr["notes"].append(nn)
            pc(doc, pr)
            continue
        if st.startswith("### "):
            hh(doc, md_clean(st[4:]), 3)
            continue
        if st.startswith("## "):
            hh(doc, md_clean(st[3:]), 2)
            continue
        if st.startswith("> "):
            pp(doc, md_clean(st[2:]), italic=True, color=GY)
            continue
        if is_bul(st):
            bl(doc, md_clean(strip_bul(st)))
            continue
        pp(doc, md_clean(st))
    if in_t[0]:
        flush()


SLUG_FOLDER = {
    "00-Pitch-Executivo": "00-Indice-Mestre",
    "01-Comecando": "01-Comecando",
    "02-Admin-Config": "02-Admin-Config",
    "03-Admin-Kanban": "03-Admin-Kanban",
    "04-Admin-IA": "04-Admin-IA",
    "05-Admin-Indicadores": "05-Admin-Indicadores",
    "06-Admin-Atividades": "06-Admin-Atividades",
    "07-Correspondente": "07-Correspondente",
    "08-Corretor": "08-Corretor",
    "09-Garantia": "09-Garantia",
    "10-Vistoria": "10-Vistoria",
    "11-Contrato": "11-Contrato",
    "12-Fluxos-Fim-a-Fim": "12-Fluxos-Fim-a-Fim",
    "13-Glossario-FAQ": "13-Glossario-FAQ",
}


def build(ch, ak, asub, alabel, acolor):
    mp = os.path.join(CONTENT_DIR, ch["md_file"])
    if not os.path.exists(mp):
        return None
    parsed = parse_chapter(mp)
    aud = ch[ak]
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)
    cover(doc, ch, alabel, acolor)
    hh(doc, "Por que este modulo importa", 1)
    pp(doc, aud["headline"], sz=14, bold=True, color=GR)
    pp(doc, aud["hook"], sz=12, italic=True)
    hh(doc, "Ao final deste modulo, voce sera capaz de:", 2)
    for o in aud["objectives"]:
        bl(doc, o)
    doc.add_page_break()
    hh(doc, ch["title"], 1)
    if parsed.get("intro"):
        render(doc, parsed["intro"])
    for sec in parsed["sections"]:
        hh(doc, sec["heading"], 2)
        render(doc, sec["body"])
    doc.add_page_break()
    hh(doc, "Recap - o que voce leva deste modulo", 1)
    for i, ti in enumerate(aud["takeaways"], 1):
        pp(doc, "#" + ("%02d" % i) + "  " + ti, sz=12, bold=True,
           color=(GR if i % 2 else GD))
    if aud.get("next_module"):
        doc.add_paragraph()
        hh(doc, "Proximo passo", 2)
        pp(doc, aud["next_module"], sz=12, bold=True, color=GD)
    doc.add_paragraph()
    pp(doc, "OmniMob - Padrao Diamond - Manual oficial",
       sz=9, italic=True, color=GY, align=WD_ALIGN_PARAGRAPH.CENTER)
    folder = os.path.join(OUT_BASE, SLUG_FOLDER[ch["slug"]], asub)
    os.makedirs(folder, exist_ok=True)
    fn = "Manual - " + ch["slug"] + " (" + ak + ").docx"
    op = os.path.join(folder, fn)
    doc.save(op)
    dp = os.path.join(DIST_DIR, fn)
    shutil.copy2(op, dp)
    return op


AUD = [("decisor", "Para o Decisor", "Para o Decisor", GD),
       ("operador", "Para o Operador", "Para o Operador", GR)]


def main():
    tot = 0
    for ch in CHAPTERS:
        print(">>> Modulo " + ("%02d" % ch["module_num"]) + " - " + ch["title"])
        for ak, asub, alabel, acolor in AUD:
            try:
                out = build(ch, ak, asub, alabel, acolor)
                if out:
                    print("  [OK] " + alabel + " -> " + os.path.basename(out))
                    tot += 1
            except Exception as e:
                print("  [ERRO] " + alabel + ": " + str(e))
                import traceback
                traceback.print_exc()
    print("Total: " + str(tot))


if __name__ == "__main__":
    main()
